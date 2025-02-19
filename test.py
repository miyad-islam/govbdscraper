import json
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

# Load the extracted information from the JSON file
with open("Output/rda.gov.bd/extracted_info.json", "r", encoding="utf-8") as f:
    extracted_data = json.load(f)

# Ask the user for the sequence number
seq_number = int(input("Enter the sequence number to extract text: "))

# Find the corresponding entry in the extracted data
selected_entry = None
for entry in extracted_data:
    if entry["seq"] == seq_number:
        selected_entry = entry
        break

if not selected_entry:
    print(f"No entry found with sequence number {seq_number}.")
    exit()

# Extract the text and URL
url = selected_entry["url"]
text = selected_entry["text"]

# Create a new Word document
doc = Document()

# Set a Bangla-supported font (e.g., "Kalpurush" or "Siyam Rupali")
bangla_font_name = "Kalpurush"  # Change this to the Bangla font you want to use

# Add the URL as a header
header = doc.sections[0].header
header_paragraph = header.paragraphs[0]
header_paragraph.text = f"Source URL: {url}"
header_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Set the font for the header
for run in header_paragraph.runs:
    run.font.name = bangla_font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), bangla_font_name)  # type: ignore

# Add the text to the document
paragraphs = text.split("\n")

for paragraph in paragraphs:
    if paragraph.strip():  # Skip empty lines
        p = doc.add_paragraph(paragraph.strip())
        p.style = "Normal"  # Use "Normal" for body text

        # Set the font for the paragraph
        for run in p.runs:
            run.font.name = bangla_font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), bangla_font_name)  # type: ignore

        # Check if the paragraph is a heading
        if paragraph.strip().isupper() or paragraph.strip().startswith(("শিরোনাম", "অধ্যায়", "ভাগ")):
            p.style = "Heading 1"  # Use "Heading 1" for headings

# Save the Word document
output_file = f"extracted_text_seq_{seq_number}.docx"
doc.save(output_file)

print(f"Text for sequence number {seq_number} has been written to {output_file}.")